"""
Multi-Format Document Processor
=================================

Handles extraction of text, tables, figures, equations, and references
from various document formats (PDF, DOCX, LaTeX, HTML).

Supports:
- PDF: Using existing PDFProcessor
- DOCX: Using python-docx
- LaTeX: Using pylatexenc (text extraction)
- HTML: Using BeautifulSoup4
"""

from typing import Dict, List, Optional, Tuple
from pathlib import Path
import re


class DocumentProcessor:
    """
    Unified document processor supporting multiple formats.

    Provides a consistent interface for extracting content from:
    - PDF files (.pdf)
    - Word documents (.docx)
    - LaTeX files (.tex)
    - HTML files (.html)

    All formats return a standardized structure:
    {
        'success': bool,
        'format': str,
        'pages': List[Dict],  # For PDF
        'sections': Dict[str, str],
        'tables': List[Dict],
        'figures': List[Dict],
        'equations': List[str],
        'references': List[str],
        'full_text': str,
        'metadata': Dict
    }
    """

    def __init__(self):
        """Initialize document processor with format handlers"""
        from rag_system.pdf_processor import PDFProcessor

        self.pdf_processor = PDFProcessor()
        self.supported_formats = ['pdf', 'docx', 'tex', 'html']

    def process_document(
        self,
        file_path: str,
        file_format: Optional[str] = None
    ) -> Dict:
        """
        Process a document and extract content.

        Args:
            file_path: Path to the document file
            file_format: Format hint ('pdf', 'docx', 'tex', 'html').
                        Auto-detected from extension if not provided.

        Returns:
            Dictionary with extracted content

        Example:
            processor = DocumentProcessor()
            result = processor.process_document('paper.pdf')
            print(result['full_text'])
        """
        # Detect format from file extension
        if file_format is None:
            file_format = Path(file_path).suffix.lstrip('.').lower()

        # Validate format
        if file_format not in self.supported_formats:
            return {
                'success': False,
                'format': file_format,
                'error': f'Unsupported format: {file_format}. Supported: {self.supported_formats}'
            }

        # Route to appropriate handler
        try:
            if file_format == 'pdf':
                return self._process_pdf(file_path)
            elif file_format == 'docx':
                return self._process_docx(file_path)
            elif file_format == 'tex':
                return self._process_latex(file_path)
            elif file_format == 'html':
                return self._process_html(file_path)
            else:
                return {
                    'success': False,
                    'format': file_format,
                    'error': f'Handler not implemented for: {file_format}'
                }

        except Exception as e:
            return {
                'success': False,
                'format': file_format,
                'error': f'Processing failed: {str(e)}'
            }

    def _process_pdf(self, file_path: str) -> Dict:
        """
        Process PDF file using existing PDFProcessor.

        Args:
            file_path: Path to PDF file

        Returns:
            Standardized content dictionary
        """
        # Use existing PDF processor
        extraction_result = self.pdf_processor.extract_text_by_sections(file_path)

        # Convert to standardized format
        sections = extraction_result.get('sections', {})
        pages = extraction_result.get('pages', [])

        # Extract full text
        full_text = '\n'.join([p.get('text', '') for p in pages])

        # Try to extract tables and figures from text
        tables = self._extract_tables_from_text(full_text)
        figures = self._extract_figures_from_text(full_text)
        equations = self._extract_equations_from_text(full_text)
        references = self._extract_references_from_sections(sections)

        return {
            'success': True,
            'format': 'pdf',
            'pages': pages,
            'sections': sections,
            'tables': tables,
            'figures': figures,
            'equations': equations,
            'references': references,
            'full_text': full_text,
            'metadata': {
                'page_count': len(pages),
                'section_count': len(sections),
                'file_path': file_path
            }
        }

    def _process_docx(self, file_path: str) -> Dict:
        """
        Process DOCX file using python-docx.

        Args:
            file_path: Path to DOCX file

        Returns:
            Standardized content dictionary
        """
        try:
            from docx import Document
        except ImportError:
            return {
                'success': False,
                'format': 'docx',
                'error': 'python-docx not installed. Install with: pip install python-docx'
            }

        try:
            doc = Document(file_path)

            # Extract text by paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = '\n'.join(paragraphs)

            # Try to identify sections based on headings
            sections = self._identify_sections_from_paragraphs(doc.paragraphs)

            # Extract tables
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)

                tables.append({
                    'table_number': i + 1,
                    'caption': f'Table {i + 1}',
                    'data': table_data,
                    'rows': len(table.rows),
                    'cols': len(table.rows[0].cells) if table.rows else 0
                })

            # Extract figures (limited - DOCX doesn't provide easy figure extraction)
            figures = self._extract_figures_from_text(full_text)

            # Extract equations and references
            equations = self._extract_equations_from_text(full_text)
            references = self._extract_references_from_sections(sections)

            return {
                'success': True,
                'format': 'docx',
                'pages': [{'page_number': 1, 'text': full_text}],  # DOCX doesn't have pages
                'sections': sections,
                'tables': tables,
                'figures': figures,
                'equations': equations,
                'references': references,
                'full_text': full_text,
                'metadata': {
                    'paragraph_count': len(paragraphs),
                    'table_count': len(tables),
                    'section_count': len(sections),
                    'file_path': file_path
                }
            }

        except Exception as e:
            return {
                'success': False,
                'format': 'docx',
                'error': f'DOCX processing failed: {str(e)}'
            }

    def _process_latex(self, file_path: str) -> Dict:
        """
        Process LaTeX file.

        Args:
            file_path: Path to .tex file

        Returns:
            Standardized content dictionary
        """
        try:
            # Read LaTeX file
            with open(file_path, 'r', encoding='utf-8') as f:
                latex_content = f.read()

            # Extract sections using regex
            sections = self._extract_latex_sections(latex_content)

            # Extract text (remove LaTeX commands - simplified)
            full_text = self._extract_text_from_latex(latex_content)

            # Extract specific LaTeX elements
            equations = self._extract_latex_equations(latex_content)
            tables = self._extract_latex_tables(latex_content)
            figures = self._extract_latex_figures(latex_content)
            references = self._extract_latex_references(latex_content)

            return {
                'success': True,
                'format': 'tex',
                'pages': [{'page_number': 1, 'text': full_text}],
                'sections': sections,
                'tables': tables,
                'figures': figures,
                'equations': equations,
                'references': references,
                'full_text': full_text,
                'metadata': {
                    'section_count': len(sections),
                    'equation_count': len(equations),
                    'file_path': file_path
                }
            }

        except Exception as e:
            return {
                'success': False,
                'format': 'tex',
                'error': f'LaTeX processing failed: {str(e)}'
            }

    def _process_html(self, file_path: str) -> Dict:
        """
        Process HTML file using BeautifulSoup.

        Args:
            file_path: Path to HTML file

        Returns:
            Standardized content dictionary
        """
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            return {
                'success': False,
                'format': 'html',
                'error': 'beautifulsoup4 not installed. Install with: pip install beautifulsoup4'
            }

        try:
            # Read HTML file
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            # Extract sections from headings
            sections = {}
            current_section = 'introduction'
            current_content = []

            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'div']):
                if element.name in ['h1', 'h2', 'h3']:
                    # Save previous section
                    if current_content:
                        sections[current_section] = '\n'.join(current_content)

                    # Start new section
                    current_section = element.get_text().strip().lower()
                    current_content = []
                elif element.name in ['p', 'div']:
                    text = element.get_text().strip()
                    if text:
                        current_content.append(text)

            # Save last section
            if current_content:
                sections[current_section] = '\n'.join(current_content)

            # Extract full text
            full_text = soup.get_text(separator='\n', strip=True)

            # Extract tables
            tables = []
            for i, table in enumerate(soup.find_all('table')):
                table_data = []
                for row in table.find_all('tr'):
                    cells = row.find_all(['td', 'th'])
                    row_data = [cell.get_text().strip() for cell in cells]
                    table_data.append(row_data)

                tables.append({
                    'table_number': i + 1,
                    'caption': f'Table {i + 1}',
                    'data': table_data
                })

            # Extract figures
            figures = []
            for i, img in enumerate(soup.find_all('img')):
                figures.append({
                    'figure_number': i + 1,
                    'src': img.get('src', ''),
                    'alt': img.get('alt', ''),
                    'caption': img.get('title', f'Figure {i + 1}')
                })

            # Extract equations and references
            equations = self._extract_equations_from_text(full_text)
            references = self._extract_references_from_sections(sections)

            return {
                'success': True,
                'format': 'html',
                'pages': [{'page_number': 1, 'text': full_text}],
                'sections': sections,
                'tables': tables,
                'figures': figures,
                'equations': equations,
                'references': references,
                'full_text': full_text,
                'metadata': {
                    'section_count': len(sections),
                    'table_count': len(tables),
                    'figure_count': len(figures),
                    'file_path': file_path
                }
            }

        except Exception as e:
            return {
                'success': False,
                'format': 'html',
                'error': f'HTML processing failed: {str(e)}'
            }

    # Helper methods for content extraction

    def _identify_sections_from_paragraphs(self, paragraphs) -> Dict[str, str]:
        """Identify sections from DOCX paragraphs based on style"""
        sections = {}
        current_section = 'introduction'
        current_content = []

        section_keywords = {
            'abstract': ['abstract'],
            'introduction': ['introduction'],
            'literature': ['literature', 'related work', 'background'],
            'methodology': ['methodology', 'methods', 'approach'],
            'results': ['results', 'experiments'],
            'discussion': ['discussion', 'analysis'],
            'conclusion': ['conclusion', 'future work'],
            'references': ['references', 'bibliography']
        }

        for para in paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a section heading
            is_heading = False
            for section_name, keywords in section_keywords.items():
                if any(keyword in text.lower() for keyword in keywords) and len(text) < 100:
                    # Save previous section
                    if current_content:
                        sections[current_section] = '\n'.join(current_content)

                    current_section = section_name
                    current_content = []
                    is_heading = True
                    break

            if not is_heading:
                current_content.append(text)

        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)

        return sections

    def _extract_text_from_latex(self, latex_content: str) -> str:
        """Extract plain text from LaTeX (remove commands)"""
        # Remove comments
        text = re.sub(r'%.*$', '', latex_content, flags=re.MULTILINE)

        # Remove common commands (simplified)
        text = re.sub(r'\\[a-zA-Z]+(\[.*?\])?(\{.*?\})?', '', text)

        # Remove math environments (keep content)
        text = re.sub(r'\$\$(.*?)\$\$', r'\1', text, flags=re.DOTALL)
        text = re.sub(r'\$(.*?)\$', r'\1', text)

        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text).strip()

        return text

    def _extract_latex_sections(self, latex_content: str) -> Dict[str, str]:
        """Extract sections from LaTeX"""
        sections = {}

        # Find section commands
        section_pattern = r'\\section\{([^}]+)\}(.*?)(?=\\section|\\end\{document\}|$)'
        matches = re.findall(section_pattern, latex_content, re.DOTALL)

        for section_name, section_content in matches:
            clean_name = section_name.strip().lower()
            clean_content = self._extract_text_from_latex(section_content)
            sections[clean_name] = clean_content

        return sections

    def _extract_latex_equations(self, latex_content: str) -> List[str]:
        """Extract equations from LaTeX"""
        equations = []

        # Find display math
        display_math = re.findall(r'\$\$(.*?)\$\$', latex_content, re.DOTALL)
        equations.extend([eq.strip() for eq in display_math])

        # Find equation environments
        eq_envs = re.findall(r'\\begin\{equation\}(.*?)\\end\{equation\}', latex_content, re.DOTALL)
        equations.extend([eq.strip() for eq in eq_envs])

        return equations

    def _extract_latex_tables(self, latex_content: str) -> List[Dict]:
        """Extract table captions from LaTeX"""
        tables = []

        # Find table environments
        table_pattern = r'\\begin\{table\}(.*?)\\end\{table\}'
        table_matches = re.findall(table_pattern, latex_content, re.DOTALL)

        for i, table_content in enumerate(table_matches):
            # Try to find caption
            caption_match = re.search(r'\\caption\{([^}]+)\}', table_content)
            caption = caption_match.group(1) if caption_match else f'Table {i + 1}'

            tables.append({
                'table_number': i + 1,
                'caption': caption,
                'content': table_content
            })

        return tables

    def _extract_latex_figures(self, latex_content: str) -> List[Dict]:
        """Extract figure captions from LaTeX"""
        figures = []

        # Find figure environments
        figure_pattern = r'\\begin\{figure\}(.*?)\\end\{figure\}'
        figure_matches = re.findall(figure_pattern, latex_content, re.DOTALL)

        for i, figure_content in enumerate(figure_matches):
            # Try to find caption
            caption_match = re.search(r'\\caption\{([^}]+)\}', figure_content)
            caption = caption_match.group(1) if caption_match else f'Figure {i + 1}'

            figures.append({
                'figure_number': i + 1,
                'caption': caption,
                'content': figure_content
            })

        return figures

    def _extract_latex_references(self, latex_content: str) -> List[str]:
        """Extract bibliography from LaTeX"""
        references = []

        # Find bibitem entries
        bibitem_pattern = r'\\bibitem\{[^}]+\}(.*?)(?=\\bibitem|\\end\{thebibliography\}|$)'
        bibitem_matches = re.findall(bibitem_pattern, latex_content, re.DOTALL)

        for ref in bibitem_matches:
            clean_ref = self._extract_text_from_latex(ref).strip()
            if clean_ref:
                references.append(clean_ref)

        return references

    def _extract_tables_from_text(self, text: str) -> List[Dict]:
        """Extract table references from text"""
        tables = []
        table_pattern = r'Table\s+(\d+)[:\s]+([^\n]{10,200})'
        matches = re.findall(table_pattern, text, re.IGNORECASE)

        for table_num, caption in matches:
            tables.append({
                'table_number': int(table_num),
                'caption': caption.strip()
            })

        return tables

    def _extract_figures_from_text(self, text: str) -> List[Dict]:
        """Extract figure references from text"""
        figures = []
        figure_pattern = r'(?:Figure|Fig\.)\s+(\d+)[:\s]+([^\n]{10,200})'
        matches = re.findall(figure_pattern, text, re.IGNORECASE)

        for fig_num, caption in matches:
            figures.append({
                'figure_number': int(fig_num),
                'caption': caption.strip()
            })

        return figures

    def _extract_equations_from_text(self, text: str) -> List[str]:
        """Extract equation references from text"""
        equations = []

        # Find inline math (simplified)
        inline_math = re.findall(r'\$([^$]+)\$', text)
        equations.extend([eq.strip() for eq in inline_math if len(eq.strip()) > 5])

        # Find equation indicators
        eq_indicators = re.findall(r'(?:Equation|Eq\.)\s+\(?\d+\)?[:\s]+([^\n]{10,200})', text, re.IGNORECASE)
        equations.extend([eq.strip() for eq in eq_indicators])

        return list(set(equations))[:50]  # Limit to 50 unique equations

    def _extract_references_from_sections(self, sections: Dict[str, str]) -> List[str]:
        """Extract individual references from references section"""
        references = []

        # Look for references/bibliography section
        ref_section = None
        for key in ['references', 'bibliography', 'References', 'Bibliography', 'REFERENCES']:
            if key in sections:
                ref_section = sections[key]
                break

        if ref_section:
            # Split by common reference patterns
            # Pattern 1: [1] Author, Title...
            pattern1 = re.split(r'\[\d+\]', ref_section)
            if len(pattern1) > 1:
                references.extend([ref.strip() for ref in pattern1 if len(ref.strip()) > 20])

            # Pattern 2: 1. Author, Title...
            if not references:
                pattern2 = re.split(r'^\d+\.', ref_section, flags=re.MULTILINE)
                references.extend([ref.strip() for ref in pattern2 if len(ref.strip()) > 20])

        return references[:100]  # Limit to 100 references


if __name__ == "__main__":
    # Test the DocumentProcessor
    import sys
    sys.path.insert(0, str(Path(__file__).parent.parent))

    print("Testing DocumentProcessor...")

    try:
        processor = DocumentProcessor()
        print(f"✓ Supported formats: {processor.supported_formats}")

        # Test format detection
        test_files = [
            'paper.pdf',
            'document.docx',
            'article.tex',
            'page.html'
        ]

        for file_path in test_files:
            detected_format = Path(file_path).suffix.lstrip('.').lower()
            print(f"✓ {file_path} -> {detected_format}")

        print("\n✅ DocumentProcessor initialized successfully!")
        print("\nNote: To test actual file processing, provide real document files.")

    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
